from pathlib import Path
from typing import Dict, List


SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx 未安装，无法读取 .docx 文档") from exc

    document = Document(path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf 未安装，无法读取 .pdf 文档") from exc

    reader = PdfReader(str(path))
    pages: List[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)

    return "\n\n".join(pages).strip()


def _read_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _read_txt(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    return ""


def load_documents(docs_dir: str) -> List[Dict[str, str]]:
    """
    加载指定目录下的所有文档内容。
    支持 .txt、.docx 和 .pdf。
    """
    docs_path = Path(docs_dir)
    if not docs_path.exists():
        return []

    documents = []
    if docs_path.is_file():
        content = _read_document(docs_path)
        if content:
            documents.append({"content": content, "source": docs_path.name})
    else:
        for file in sorted(docs_path.iterdir()):
            if not file.is_file() or file.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            content = _read_document(file)
            if content:
                documents.append({"content": content, "source": file.name})

    return documents
